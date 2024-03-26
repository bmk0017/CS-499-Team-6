import string
import docx
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
import re

# vvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvvv
# *****************************************************

# TO DO:
# Test to make sure everything smooth/wait for actual question objects to be generated
# add additional functionality as needed

# *****************************************************
# ^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^


# word doc generator class that can be imported if needed
class QuizDocGenerator:
    # all these add functions are private atm

    # add multiple choice question, given document, question object (applies for T/F as well)
    def __add_MC_question(self, document, qobj, key):
        question = document.add_paragraph(style='List Number')
        question.add_run(remove_html_tags(qobj.question_content) + '\n', style='userstyle')
        question.paragraph_format.keep_together = True
        i = 0
        for id, value in qobj.choices.items():
            if key and id == qobj.correct_choice:
                question.add_run(string.ascii_lowercase[i] + ')\t\t' + value + '\n', style='userstyle').font.color.rgb = RGBColor(0, 255,
                                                                                                               0)
            else:
                question.add_run(string.ascii_lowercase[i] + ')\t\t' + value + '\n', style ='userstyle')
            i += 1

    # adds true/false question
    def __add_TF_question(self, document, qobj, key):
        question = document.add_paragraph(style='List Number')
        question.add_run(remove_html_tags(qobj.question_content) + '\n', style='userstyle')
        question.paragraph_format.keep_together = True
        i = 0
        for id, value in qobj.choices.items():
            if key and id == qobj.correct_choice:
                question.add_run(string.ascii_lowercase[i] + ')\t\t' + value + '\n', style='userstyle').font.color.rgb = RGBColor(0, 255,
                                                                                                               0)
            elif not key:
                question.add_run(string.ascii_lowercase[i] + ')\t\t' + value + '\n', style='userstyle')
            i += 1

    # Adds multiple answer question
    def __add_MA_question(self, document, qobj, key):
        question = document.add_paragraph(style='List Number')
        question.add_run(remove_html_tags(qobj.question_content) + '\n', style='userstyle')
        question.paragraph_format.keep_together = True
        i = 0
        for id, value in qobj.choices.items():
            if key and id in qobj.correct_choices:
                question.add_run(string.ascii_lowercase[i] + ')\t\t' + value + '\n', style='userstyle').font.color.rgb = RGBColor(0, 255,
                                                                                                               0)
            else:
                question.add_run(string.ascii_lowercase[i] + ')\t\t' + value + '\n', style='userstyle')
            i += 1

    # add fill in the blank question (also applies for multiple fill in the blank); I'm assuming that the prompt would have the blank in it so just printing prompt
    def __add_FnB_question(self, document, qobj, key):
        question = document.add_paragraph(style='List Number')
        question.add_run(remove_html_tags(qobj.question_content) + '\n', style='userstyle')
        question.paragraph_format.keep_together = True
        if key:
            question.add_run(str(qobj.correct_choices), style='userstyle').font.color.rgb = RGBColor(0, 255, 0)
            question.add_run('\n')

    # adds question format for when ther is a separate question and (user inputed/typed/written) answer field/blank
    def __add_Short_question(self, document, qobj, key):
        question = document.add_paragraph(style='List Number')
        question.add_run(remove_html_tags(qobj.question_content) + '\n', style='userstyle')
        question.paragraph_format.keep_together = True
        if key:
            question.add_run(qobj.correct_feedback, style='userstyle').font.color.rgb = RGBColor(0, 255, 0)
        question.add_run('\n\n\n')

    # adds numeric question
    def __add_Num_question(self, document, qobj, key):
        question = document.add_paragraph(style='List Number')
        question.add_run(remove_html_tags(qobj.question_content) + '\n', style='userstyle')
        question.paragraph_format.keep_together = True
        if key:
            question.add_run("Answer Range:\t" + str(qobj.answer_ranges), style='userstyle').font.color.rgb = RGBColor(0, 255, 0)
        question.add_run('\n\n\n')

    # adds formula question
    def __add_Form_question(self, document, qobj, key):
        # find chosen id
        qID = None
        for i in range(0, len(qobj.possible_questions)):
            if qobj.possible_questions[i].id == qobj.chosenID:
                qID = i

        question = document.add_paragraph(style='List Number')
        answkey = list(qobj.possible_questions[qID].variable_name_value.keys())[0]
        txt = remove_html_tags(qobj.question_content)
        txt = txt.replace('[' + str(answkey) + ']', qobj.possible_questions[qID].variable_name_value[answkey])
        question.add_run(txt + '\n', style='userstyle')
        question.paragraph_format.keep_together = True
        if key:
            question.add_run(qobj.possible_questions[qID].answers, style='userstyle').font.color.rgb = RGBColor(0, 255, 0)
        question.add_run('\n\n\n')

    # adds question format for long response questions (applicable for file upload questions?); each of these has its own page
    def __add_Essay_question(self, document, qobj, key):
        question = document.add_paragraph(style='List Number')
        question.add_run(remove_html_tags(qobj.question_content) + '\n', style='userstyle')
        question.paragraph_format.page_break_before = True
        if key:
            question.add_run(qobj.feedback, style='userstyle').font.color.rgb = RGBColor(0, 255, 0)
        document.add_page_break()

    # adds text field
    def __add_Text_field(self, document, qobj):
        text = document.add_paragraph()
        text.add_run(remove_html_tags(qobj.question_content) + '\n', style='userstyle')
        text.paragraph_format.keep_together = True

    # adds matching question
    def __add_Match_question(self, document, qobj, key):
        if key:
            question = document.add_paragraph(style='List Number')
            question.add_run(remove_html_tags(qobj.question_content) + '\n', style='userstyle')
            question.paragraph_format.keep_together = True
            for id1, id2 in qobj.correct_choices.items():
                question.add_run(
                    qobj.left_choices[id1] + ':\t-->\t' + qobj.right_choices[id2] + '\n').font.color.rgb = RGBColor(0,
                                                                                                                    255,
                                                                                                                    0)
        else:
            outtable = document.add_table(rows=1, cols=1).cell(0, 0)
            table = outtable.add_table(rows=1, cols=4)
            row = table.rows[0].cells
            row[0].add_paragraph(style='List Number')
            row[0].add_paragraph().add_run(remove_html_tags(qobj.question_content), style='userstyle')
            row[0].merge(row[3])
            num_rows = max(len(qobj.right_choices), len(qobj.left_choices))
            index = 1
            for i in range(0, num_rows):
                temp_row = table.add_row().cells
            for id in qobj.right_choices:
                row = table.rows[index].cells
                row[3].add_paragraph().add_run(qobj.right_choices[id], style='userstyle')
                index += 1
            index = 1
            for id in qobj.left_choices:
                row = table.rows[index].cells
                row[0].add_paragraph().add_run(qobj.left_choices[id], style='userstyle')
                index += 1
            table.add_row()

    # adds multiple dropdown question
    def __add_Mdrop_question(self, document, qobj, key):
        question = document.add_paragraph(style='List Number')
        question.add_run(remove_html_tags(qobj.question_content) + '\n', style='userstyle')
        question.paragraph_format.keep_together = True
        if key:
            for name, id in qobj.correct_choices.items():
                tempdict = qobj.choices.get(name)
                question.add_run(name + ':\t' + tempdict.get(id) + '\n', style='userstyle').font.color.rgb = RGBColor(0, 255, 0)
        else:
            for dname in qobj.choices:
                i = 0
                question.add_run(dname + '\n')
                for value in qobj.choices.get(dname).values():
                    question.add_run('\t' + string.ascii_lowercase[i] + ')\t\t' + value + '\n', style='userstyle')
                    i += 1

    # read questions from array from quiz object/class this is the only public function atm
    def Quiz_to_Doc(self, quiz, key=False, fontsize = 12, fonttype = 'Times New Roman'):

        #font stuff
        document = docx.Document()
        font_styles = document.styles
        font_charstyle = font_styles.add_style('userstyle', WD_STYLE_TYPE.CHARACTER)
        font_object = font_charstyle.font
        font_object.size = Pt(fontsize)
        font_object.name = fonttype


        document.add_heading(quiz.title, 0)
        document.add_heading(quiz.description, 1)
        for i in range(0, len(quiz.questions)):
            qtype = quiz.questions[i].question_type
            if qtype == "multiple_choice_question":
                self.__add_MC_question(document, quiz.questions[i], key)
            elif qtype == "true_false_question":
                self.__add_TF_question(document, quiz.questions[i], key)
            elif qtype == "short_answer_question" or qtype == "fill_in_multiple_blanks_question":
                self.__add_FnB_question(document, quiz.questions[i], key)
            elif qtype == "multiple_answers_question":
                self.__add_MA_question(document, quiz.questions[i], key)
            elif qtype == "matching_question":
                self.__add_Match_question(document, quiz.questions[i], key)
            elif qtype == "numerical_question":
                self.__add_Num_question(document, quiz.questions[i], key)
            elif qtype == "calculated_question":
                self.__add_Form_question(document, quiz.questions[i], key)
            elif qtype == "essay_question" or qtype == "file_upload_question":
                self.__add_Essay_question(document, quiz.questions[i], key)
            elif qtype == "text_only_question":
                self.__add_Text_field(document, quiz.questions[i])
            elif qtype == "multiple_dropdowns_question":
                self.__add_Mdrop_question(document, quiz.questions[i], key)

        if key:
            document.save(quiz.title + '_' + quiz.id + '_key.docx')
        else:
            document.save(quiz.title + '_' + quiz.id + '.docx')

def remove_html_tags(text):
    """Remove html tags from a string"""
    clean = re.compile('<.*?>')
    return re.sub(clean, '', text).replace('&nbsp', '')
